#!/usr/bin/env python3
"""Generate Contract de Cesiune Exclusiva a Drepturilor Patrimoniale de Autor - Word Document"""

import os
import sys

# Try to import docx, install if not available
try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("python-docx not found, attempting to install...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml


def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, top={"sz": 6, "color": "000000", "val": "single"}, ...)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", 4)}" w:space="0" w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_table_borders(table):
    """Set borders for all cells in table"""
    border_args = {
        "top": {"sz": 4, "color": "000000", "val": "single"},
        "bottom": {"sz": 4, "color": "000000", "val": "single"},
        "start": {"sz": 4, "color": "000000", "val": "single"},
        "end": {"sz": 4, "color": "000000", "val": "single"},
    }
    # Set table-level borders via XML
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_paragraph(doc, text, font_size=11, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=6, font_name='Arial', first_line_indent=None):
    """Add a paragraph with specified formatting"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    # Set font for complex script (Romanian chars)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:cs'), font_name)
    return p


def add_mixed_paragraph(doc, parts, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        space_before=0, space_after=6, font_name='Arial'):
    """Add a paragraph with mixed formatting. parts is list of (text, font_size, bold) tuples"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    for text, font_size, bold in parts:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="{font_name}"/>')
            rPr.append(rFonts)
        else:
            rFonts.set(qn('w:cs'), font_name)
    return p


def add_chapter_heading(doc, text):
    """Add a chapter heading"""
    return add_paragraph(doc, text, font_size=13, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        space_before=18, space_after=10)


def add_normal(doc, text, **kwargs):
    """Add normal text paragraph"""
    return add_paragraph(doc, text, font_size=11, bold=False, **kwargs)


def add_bold(doc, text, **kwargs):
    """Add bold text paragraph"""
    return add_paragraph(doc, text, font_size=11, bold=True, **kwargs)


def main():
    output_dir = r"D:\SoundFree\Acte"
    output_file = os.path.join(output_dir, "Contract_Cesiune_Drepturi_Autor_EGI.docx")

    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Set A4 page size
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ===== TITLE =====
    add_paragraph(doc, "CONTRACT DE CESIUNE EXCLUSIV\u0102",
                  font_size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=12, space_after=2)
    add_paragraph(doc, "A DREPTURILOR PATRIMONIALE DE AUTOR",
                  font_size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=2, space_after=16)

    add_paragraph(doc, "Nr. _____ / Data: ___.___.______",
                  font_size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=6, space_after=12)

    add_normal(doc,
        "Încheiat în temeiul Legii nr. 8/1996 privind dreptul de autor și drepturile conexe, "
        "cu modificările și completările ulterioare (inclusiv Legea nr. 74/2018), și al "
        "Directivei 2014/26/UE a Parlamentului European și a Consiliului.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

    # ===== CAPITOLUL I =====
    add_chapter_heading(doc, "CAPITOLUL I \u2013 P\u0102R\u021aILE CONTRACTANTE")

    add_bold(doc, "1.1. CESIONARUL (Entitatea de Gestiune Independent\u0103):",
             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

    fields_cesionar = [
        "Denumirea: _________________________________ S.R.L.",
        "CUI: _________________________________",
        "Nr. Reg. Com.: _________________________________",
        "Sediul: _________________________________",
        "Cont bancar: _________________________________",
        "Banca: _________________________________",
        "Reprezentant: _________________________________, în calitate de Administrator",
        "E-mail: _________________________________",
        "Telefon: _________________________________",
    ]
    for f in fields_cesionar:
        add_normal(doc, f, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)

    add_normal(doc, "", space_after=2)
    add_normal(doc,
        "Înregistrat\u0103 ca Entitate de Gestiune Independent\u0103 (EGI) la Oficiul Român "
        "pentru Drepturile de Autor (ORDA) conform art. 123^5 din Legea nr. 8/1996.",
        space_after=6)

    add_normal(doc, 'denumit în continuare \u201eEGI\u201d sau \u201eCesionarul\u201d,',
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

    add_bold(doc, "\u0218I", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=10)

    add_bold(doc, "1.2. CEDENTUL (Autorul/Compozitorul):",
             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

    fields_cedent = [
        "Numele \u0219i prenumele: _________________________________",
        "CNP: _________________________________",
        "Seria CI/BI: _____ Nr.: _______________________",
        "Domiciliul: _________________________________",
        "Cont bancar (IBAN): _________________________________",
        "Banca: _________________________________",
        "E-mail: _________________________________",
        "Telefon: _________________________________",
        "Pseudonim artistic: _________________________________ (dac\u0103 este cazul)",
    ]
    for f in fields_cedent:
        add_normal(doc, f, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)

    add_normal(doc, "", space_after=2)
    add_normal(doc, 'denumit în continuare \u201eAutorul\u201d sau \u201eCedentul\u201d.',
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

    # ===== CAPITOLUL II =====
    add_chapter_heading(doc, "CAPITOLUL II \u2013 OBIECTUL CONTRACTULUI")

    add_normal(doc,
        "2.1. Prin prezentul contract, Autorul cedeaz\u0103 în mod EXCLUSIV c\u0103tre EGI "
        "drepturile patrimoniale de autor asupra operelor muzicale enumerate în ANEXA NR. 1, "
        "care face parte integrant\u0103 din prezentul contract (denumite în continuare \u201eOperele\u201d).")

    add_normal(doc,
        "2.2. Operele sunt crea\u021bii originale ale Autorului, create anterior sau pe parcursul "
        "derul\u0103rii prezentului contract, \u0219i livrate EGI-ului în format digital "
        "(WAV, FLAC, MP3 sau alt format convenit de p\u0103r\u021bi).")

    add_mixed_paragraph(doc, [
        ("2.3. ", 11, False),
        ("IMPORTANT: ", 11, True),
        ("Prezentul contract NU are ca obiect cesiunea drepturilor patrimoniale asupra "
         "totalit\u0103\u021bii operelor viitoare ale Autorului, ci exclusiv asupra operelor "
         "identificate individual în Anexa nr. 1, care poate fi completat\u0103 periodic prin "
         "Acte Adi\u021bionale semnate de ambele p\u0103r\u021bi, conform art. 40 alin. (1) "
         "din Legea nr. 8/1996.", 11, False),
    ])

    # ===== CAPITOLUL III =====
    add_chapter_heading(doc, "CAPITOLUL III \u2013 DREPTURILE PATRIMONIALE CEDATE")

    add_normal(doc,
        "3.1. Autorul cedeaz\u0103 în mod exclusiv c\u0103tre EGI urm\u0103toarele drepturi "
        "patrimoniale asupra Operelor, conform art. 13 din Legea nr. 8/1996:")

    items_3_1 = [
        ("a) Dreptul de ", "REPRODUCERE", " a Operelor, integral sau par\u021bial, pe orice tip de suport, inclusiv digital;"),
        ("b) Dreptul de ", "DISTRIBUIRE", " a Operelor, inclusiv prin mijloace electronice \u0219i platforme digitale;"),
        ("c) Dreptul de ", "COMUNICARE PUBLIC\u0102", " a Operelor, inclusiv difuzarea ca muzic\u0103 ambientală în spa\u021bii comerciale (magazine, restaurante, hoteluri, birouri, centre comerciale, s\u0103li de a\u0219teptare etc.);"),
        ("d) Dreptul de ", "PUNERE LA DISPOZI\u021aIA PUBLICULUI", " a Operelor, astfel încât orice persoan\u0103 s\u0103 poat\u0103 avea acces la acestea din orice loc \u0219i în orice moment ales în mod individual (streaming, desc\u0103rcare);"),
        ("e) Dreptul de ", "RADIODIFUZARE", " \u0219i ", "RETRANSMITERE PRIN CABLU", " a Operelor;"),
        ("f) Dreptul de a realiza ", "OPERE DERIVATE", " (adapt\u0103ri, aranjamente, remixuri), exclusiv în scopul utiliz\u0103rii ca muzic\u0103 ambientală."),
    ]

    for item in items_3_1:
        if len(item) == 3:
            parts = [(item[0], 11, False), (item[1], 11, True), (item[2], 11, False)]
        elif len(item) == 5:
            parts = [(item[0], 11, False), (item[1], 11, True), (item[2], 11, False),
                     (item[3], 11, True), (item[4], 11, False)]
        else:
            parts = [(item[0], 11, False), (item[1], 11, True), (item[2], 11, False)]
        add_mixed_paragraph(doc, parts, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

    add_normal(doc, "", space_after=2)
    add_mixed_paragraph(doc, [
        ("3.2. Cesiunea este ", 11, False),
        ("EXCLUSIV\u0102", 11, True),
        (", ceea ce înseamn\u0103 c\u0103, pe durata prezentului contract \u0219i pentru "
         "teritoriul convenit:", 11, False),
    ])

    add_normal(doc, "- Autorul NU mai poate utiliza Operele în modalit\u0103\u021bile cedate;",
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc, "- Autorul NU mai poate transmite acelea\u0219i drepturi unei alte persoane "
              "fizice sau juridice, inclusiv unui OGC sau altui EGI.",
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

    add_normal(doc,
        "3.3. EGI are dreptul de a sublicen\u021bia drepturile cedate c\u0103tre ter\u021bi "
        "(clien\u021bi ai platformei SoundFree.ro) în scopul utiliz\u0103rii Operelor ca "
        "muzic\u0103 ambientală în spa\u021bii comerciale.")

    # ===== CAPITOLUL IV =====
    add_chapter_heading(doc, "CAPITOLUL IV \u2013 TERITORIUL \u0218I DURATA")

    add_bold(doc, "4.1. TERITORIUL:", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_normal(doc, "Cesiunea este valabil\u0103 pe teritoriul:", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc, "\u2610 României", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc, "\u2610 Uniunii Europene", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc, "\u2610 Mondial", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc, "(se bifeaz\u0103 op\u021biunea convenit\u0103 de p\u0103r\u021bi)",
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

    add_bold(doc, "4.2. DURATA:", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_normal(doc,
        "Prezentul contract se încheie pe o perioad\u0103 de _____ ani, începând de la data semn\u0103rii.")

    add_normal(doc, "4.3. La expirarea duratei, contractul:", space_after=2)
    add_normal(doc, "\u2610 Înceteaz\u0103 de drept, iar drepturile revin Autorului",
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc,
        "\u2610 Se prelunge\u0219te automat cu perioade succesive de câte 1 an, dac\u0103 "
        "niciuna dintre p\u0103r\u021bi nu notific\u0103 cealalt\u0103 parte cu cel pu\u021bin "
        "60 de zile înainte de expirare",
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

    add_normal(doc,
        "4.4. Pe durata contractului, EGI poate retrage din exploatare una sau mai multe Opere, "
        "notificând Autorul în scris, f\u0103r\u0103 a datora desp\u0103gubiri suplimentare "
        "fa\u021b\u0103 de remunera\u021bia deja achitat\u0103.")

    # ===== CAPITOLUL V =====
    add_chapter_heading(doc, "CAPITOLUL V \u2013 REMUNERA\u021aIA")

    add_normal(doc,
        "5.1. În schimbul cesiunii exclusive a drepturilor patrimoniale, EGI se oblig\u0103 "
        "s\u0103 pl\u0103teasc\u0103 Autorului o remunera\u021bie dup\u0103 cum urmeaz\u0103 "
        "(se bifeaz\u0103 varianta convenit\u0103):")

    add_mixed_paragraph(doc, [
        ("\u2610 ", 11, False),
        ("VARIANTA A \u2013 Sum\u0103 fix\u0103 per oper\u0103: ", 11, True),
        ("_______ lei (brut) per oper\u0103 muzical\u0103 livrat\u0103 \u0219i acceptat\u0103. "
         "Plata se efectueaz\u0103 în termen de 30 de zile de la acceptarea operei de c\u0103tre EGI.", 11, False),
    ], space_after=4)

    add_mixed_paragraph(doc, [
        ("\u2610 ", 11, False),
        ("VARIANTA B \u2013 Redeven\u021b\u0103 periodic\u0103: ", 11, True),
        ("_______ lei (brut) / lun\u0103 pentru întregul catalog al Autorului, pl\u0103tibil\u0103 "
         "pân\u0103 la data de 15 a lunii urm\u0103toare.", 11, False),
    ], space_after=4)

    add_mixed_paragraph(doc, [
        ("\u2610 ", 11, False),
        ("VARIANTA C \u2013 Procent din încas\u0103ri: ", 11, True),
        ("_____% din veniturile nete încasate de EGI din exploatarea Operelor Autorului, calculate "
         "trimestrial. Plata se efectueaz\u0103 în termen de 30 de zile de la încheierea "
         "fiec\u0103rui trimestru calendaristic.", 11, False),
    ], space_after=4)

    add_mixed_paragraph(doc, [
        ("\u2610 ", 11, False),
        ("VARIANTA D \u2013 Hibrid: ", 11, True),
        ("_______ lei (brut) per oper\u0103 la livrare (avans) + _____% din veniturile nete "
         "din exploatare (trimestrial).", 11, False),
    ], space_after=6)

    add_normal(doc,
        "5.2. Sumele men\u021bionate sunt brute. EGI va re\u021bine la surs\u0103 impozitul "
        "pe venit conform art. 72 din Codul fiscal (10% aplicat asupra venitului net = 60% "
        "din venitul brut), precum \u0219i contribu\u021biile sociale datorate conform "
        "legisla\u021biei în vigoare (CAS/CASS), acolo unde este cazul.")

    add_normal(doc,
        "5.3. Plata se efectueaz\u0103 prin transfer bancar în contul Autorului men\u021bionat la Cap. I.")

    add_normal(doc,
        "5.4. EGI va comunica Autorului, cel pu\u021bin o dat\u0103 pe an, o situa\u021bie "
        "detaliat\u0103 privind exploatarea Operelor \u0219i veniturile generate, conform "
        "art. 123^5 alin. (8) din Legea nr. 8/1996.")

    # ===== CAPITOLUL VI =====
    add_chapter_heading(doc, "CAPITOLUL VI \u2013 DECLARA\u021aII \u0218I GARAN\u021aII ALE AUTORULUI")

    add_normal(doc,
        "6.1. Autorul declar\u0103 \u0219i garanteaz\u0103, pe propria r\u0103spundere, c\u0103:")

    items_6_1 = [
        "a) Este UNICUL AUTOR \u0219i titular al drepturilor patrimoniale asupra Operelor enumerate în Anexa nr. 1;",
        "b) Operele sunt CREA\u021aII ORIGINALE \u0219i nu încalc\u0103 drepturile de autor sau drepturile conexe ale niciunei alte persoane;",
        "c) Operele NU sunt \u0219i NU au fost înregistrate, declarate sau încredin\u021bate spre gestiune niciunui Organism de Gestiune Colectiv\u0103 (OGC), inclusiv, dar f\u0103r\u0103 a se limita la: UCMR-ADA, CREDIDAM, UPFR, sau orice alt OGC din România sau din str\u0103in\u0103tate;",
        "d) Operele NU fac obiectul niciunui alt contract de cesiune exclusiv\u0103 sau neexclusiv\u0103 cu ter\u021be persoane fizice sau juridice;",
        "e) Nu exist\u0103 niciun litigiu, revendicare sau preten\u021bie în leg\u0103tur\u0103 cu Operele la data semn\u0103rii prezentului contract;",
        "f) În cazul în care Autorul este sau a fost membru al unui OGC, Operele cedate prin prezentul contract sunt DIFERITE de cele aflate în gestiunea OGC-ului \u0219i NU au fost niciodat\u0103 declarate sau înregistrate în repertoriul administrat de acel OGC.",
    ]
    for item in items_6_1:
        add_normal(doc, item, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

    add_normal(doc,
        "6.2. Autorul se oblig\u0103 ca, pe toat\u0103 durata prezentului contract, s\u0103 NU "
        "declare, înregistreze sau cedeze Operele c\u0103tre niciun OGC sau alt\u0103 entitate, "
        "sub sanc\u021biunea rezilierii de plin drept a contractului \u0219i a pl\u0103\u021bii "
        "de daune-interese.")

    add_normal(doc,
        "6.3. În cazul în care oricare dintre declara\u021biile de mai sus se dovede\u0219te "
        "a fi neadev\u0103rat\u0103, Autorul va desp\u0103gubi integral EGI-ul pentru toate "
        "prejudiciile suferite, inclusiv sumele pl\u0103tite de EGI c\u0103tre ter\u021bi sau "
        "ca urmare a unor sanc\u021biuni aplicate de autorit\u0103\u021bi.")

    # ===== CAPITOLUL VII =====
    add_chapter_heading(doc, "CAPITOLUL VII \u2013 OBLIGA\u021aIILE EGI (CESIONARULUI)")

    add_normal(doc, "7.1. EGI se oblig\u0103 s\u0103:")

    items_7_1 = [
        "a) Exploateze Operele cu bun\u0103-credin\u021b\u0103 \u0219i în conformitate cu destina\u021bia stabilit\u0103 (muzic\u0103 ambientală pentru spa\u021bii comerciale);",
        "b) Pl\u0103teasc\u0103 remunera\u021bia convenit\u0103 la termenele stabilite;",
        "c) Comunice Autorului, cel pu\u021bin anual, o situa\u021bie privind exploatarea Operelor \u0219i veniturile generate;",
        "d) Respecte drepturile morale ale Autorului, în special:",
        "   - Dreptul de paternitate (men\u021bionarea numelui/pseudonimului Autorului acolo unde este posibil \u0219i relevant);",
        "   - Dreptul la integritatea operei (s\u0103 nu aduc\u0103 modific\u0103ri care ar prejudicia onoarea sau reputa\u021bia Autorului), cu excep\u021bia adapt\u0103rilor convenite la art. 3.1 lit. f);",
        "e) Re\u021bin\u0103 \u0219i vireze impozitele \u0219i contribu\u021biile sociale datorate conform legisla\u021biei fiscale în vigoare;",
        "f) Informeze ORDA \u0219i orice autoritate competent\u0103 cu privire la repertoriul gestionat, conform obliga\u021biilor legale.",
    ]
    for item in items_7_1:
        add_normal(doc, item, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3)

    # ===== CAPITOLUL VIII =====
    add_chapter_heading(doc, "CAPITOLUL VIII \u2013 LIVRAREA \u0218I ACCEPTAREA OPERELOR")

    add_normal(doc, "8.1. Autorul livreaz\u0103 Operele în format digital, prin:", space_after=2)
    add_normal(doc, "\u2610 Upload pe platforma dedicat\u0103 a EGI",
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc, "\u2610 Transfer electronic (e-mail, WeTransfer, Google Drive etc.)",
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc, "\u2610 Alt mod convenit: _________________________________",
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

    add_normal(doc, "8.2. Formatul tehnic minim acceptat:", space_after=2)
    add_normal(doc, "- Audio: WAV 44.1kHz / 16bit sau superior",
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc, "- Alternativ: FLAC, AIFF",
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc, "- MP3 320kbps (doar ca format suplimentar)",
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

    add_normal(doc,
        "8.3. EGI are dreptul de a accepta sau refuza o oper\u0103 muzical\u0103 în termen "
        "de 15 zile lucr\u0103toare de la livrare. Refuzul va fi comunicat în scris, motivat. "
        "Operele refuzate NU intr\u0103 sub inciden\u021ba prezentului contract, iar drepturile "
        "asupra lor r\u0103mân ale Autorului.")

    add_normal(doc,
        "8.4. La livrarea fiec\u0103rei opere noi, p\u0103r\u021bile vor semna un Act "
        "Adi\u021bional de completare a Anexei nr. 1.")

    # ===== CAPITOLUL IX =====
    add_chapter_heading(doc, "CAPITOLUL IX \u2013 ÎNCETAREA CONTRACTULUI")

    add_normal(doc, "9.1. Prezentul contract înceteaz\u0103 prin:", space_after=4)

    items_9_1 = [
        "a) Expirarea duratei, în condi\u021biile art. 4.2 \u0219i 4.3;",
        "b) Acordul scris al ambelor p\u0103r\u021bi;",
        "c) Rezilierea de c\u0103tre oricare parte, în cazul neîndeplinirii obliga\u021biilor esen\u021biale de c\u0103tre cealalt\u0103 parte, cu o notificare prealabil\u0103 de 30 de zile \u0219i acordarea unui termen de remediere de 15 zile;",
        "d) Rezilierea de plin drept, f\u0103r\u0103 punere în întârziere \u0219i f\u0103r\u0103 interven\u021bia instan\u021bei, în cazul în care declara\u021biile Autorului de la Cap. VI se dovedesc false (pact comisoriu expres de grad IV);",
        "e) Decesul Autorului \u2013 drepturile trec la mo\u0219tenitori conform legii, care pot opta pentru continuarea sau încetarea contractului.",
    ]
    for item in items_9_1:
        add_normal(doc, item, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3)

    add_normal(doc,
        "9.2. La încetarea contractului, drepturile patrimoniale cedate revin Autorului, "
        "cu excep\u021bia:", space_after=2)
    add_normal(doc,
        "- Licen\u021belor deja acordate de EGI clien\u021bilor s\u0103i, care r\u0103mân "
        "valabile pân\u0103 la expirarea lor natural\u0103;",
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc,
        "- Sumelor deja încasate de EGI, care nu se restituie.",
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

    # ===== CAPITOLUL X =====
    add_chapter_heading(doc, "CAPITOLUL X \u2013 CONFIDEN\u021aIALITATE")

    add_normal(doc,
        "10.1. P\u0103r\u021bile se oblig\u0103 s\u0103 p\u0103streze confiden\u021bialitatea "
        "termenilor financiari ai prezentului contract.")

    add_normal(doc,
        "10.2. Obliga\u021bia de confiden\u021bialitate supravie\u021buie\u0219te încet\u0103rii "
        "contractului pe o perioad\u0103 de 2 ani.")

    # ===== CAPITOLUL XI =====
    add_chapter_heading(doc, "CAPITOLUL XI \u2013 FOR\u021aA MAJOR\u0102")

    add_normal(doc,
        "11.1. Niciuna dintre p\u0103r\u021bi nu r\u0103spunde pentru neexecutarea "
        "obliga\u021biilor contractuale dac\u0103 aceasta se datoreaz\u0103 unui eveniment "
        "de for\u021b\u0103 major\u0103, astfel cum este definit de legisla\u021bia în vigoare.")

    add_normal(doc,
        "11.2. Partea afectat\u0103 va notifica cealalt\u0103 parte în termen de 5 zile de "
        "la apari\u021bia evenimentului de for\u021b\u0103 major\u0103.")

    # ===== CAPITOLUL XII =====
    add_chapter_heading(doc, "CAPITOLUL XII \u2013 LITIGII")

    add_normal(doc,
        "12.1. Orice neîn\u021belegere decurgând din sau în leg\u0103tur\u0103 cu prezentul "
        "contract va fi solu\u021bionat\u0103 pe cale amiabil\u0103.")

    add_normal(doc,
        "12.2. În cazul în care solu\u021bionarea amiabil\u0103 nu este posibil\u0103 în termen "
        "de 30 de zile, litigiul va fi supus spre solu\u021bionare instan\u021belor "
        "judec\u0103tore\u0219ti competente de la sediul Cesionarului (EGI).")

    # ===== CAPITOLUL XIII =====
    add_chapter_heading(doc, "CAPITOLUL XIII \u2013 DISPOZI\u021aII FINALE")

    add_normal(doc,
        "13.1. Prezentul contract, împreun\u0103 cu Anexele sale, constituie acordul integral "
        "al p\u0103r\u021bilor \u0219i înlocuie\u0219te orice în\u021belegere anterioar\u0103, "
        "scris\u0103 sau verbal\u0103, cu privire la obiectul s\u0103u.")

    add_normal(doc,
        "13.2. Orice modificare a prezentului contract este valabil\u0103 numai dac\u0103 "
        "este f\u0103cut\u0103 în scris \u0219i semnat\u0103 de ambele p\u0103r\u021bi "
        "(Act Adi\u021bional).")

    add_normal(doc,
        "13.3. Dac\u0103 oricare prevedere a prezentului contract este declarat\u0103 nul\u0103 "
        "sau neaplicabil\u0103, celelalte prevederi r\u0103mân valabile.")

    add_normal(doc,
        "13.4. Prezentul contract s-a încheiat în 2 (dou\u0103) exemplare originale, câte "
        "unul pentru fiecare parte.")

    add_normal(doc,
        "13.5. Prezentul contract este guvernat de legisla\u021bia din România.")

    # ===== SEMNĂTURI =====
    add_paragraph(doc, "", space_after=6)
    add_bold(doc, "SEMN\u0102TURI:", alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=12, space_after=12)

    # Create a 2-column table for signatures
    sig_table = doc.add_table(rows=6, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left column - Cesionarul
    sig_data_left = [
        ("CESIONARUL (EGI):", True),
        ("Denumirea: _______________", False),
        ("Reprezentant: ____________", False),
        ("Semn\u0103tura: _______________", False),
        ("\u0218tampila:", False),
        ("Data: ____________________", False),
    ]
    sig_data_right = [
        ("CEDENTUL (AUTORUL):", True),
        ("Numele: _______________", False),
        ("Semn\u0103tura: ____________", False),
        ("Data: _________________", False),
        ("", False),
        ("", False),
    ]

    for i, ((left_text, left_bold), (right_text, right_bold)) in enumerate(zip(sig_data_left, sig_data_right)):
        # Left cell
        cell_left = sig_table.cell(i, 0)
        cell_left.text = ""
        p = cell_left.paragraphs[0]
        run = p.add_run(left_text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = left_bold

        # Right cell
        cell_right = sig_table.cell(i, 1)
        cell_right.text = ""
        p = cell_right.paragraphs[0]
        run = p.add_run(right_text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = right_bold

    # ===== PAGE BREAK =====
    doc.add_page_break()

    # ===== ANEXA NR. 1 =====
    add_paragraph(doc, "ANEXA NR. 1 \u2013 LISTA OPERELOR MUZICALE CEDATE",
                  font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=12, space_after=4)
    add_paragraph(doc, "(parte integrant\u0103 a Contractului nr. _____)",
                  font_size=11, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=2, space_after=16)

    # Create the table
    headers = ["Nr.", "Titlul operei", "Durata\n(mm:ss)", "Gen muzical", "Data livrare",
               "Remunera\u021bie\n(lei brut)"]
    table = doc.add_table(rows=11, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Set column widths
    widths = [Cm(1.2), Cm(5.0), Cm(2.2), Cm(3.0), Cm(2.5), Cm(2.5)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        # Gray background for header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows (empty, numbered 1-10)
    for row_idx in range(1, 11):
        cell = table.cell(row_idx, 0)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(row_idx))
        run.font.name = 'Arial'
        run.font.size = Pt(10)

        # Set empty cells with proper font
        for col_idx in range(1, 6):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run("")
            run.font.name = 'Arial'
            run.font.size = Pt(10)

    # Add spacing after table
    add_paragraph(doc, "", space_after=12)

    # Declaration below table
    add_normal(doc,
        "Autorul declar\u0103 pe propria r\u0103spundere c\u0103 operele enumerate mai sus:",
        space_before=10, space_after=4)

    add_normal(doc,
        "\u2610 NU sunt înregistrate la niciun OGC (UCMR-ADA, CREDIDAM, UPFR sau altul)",
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc,
        "\u2610 NU fac obiectul niciunui alt contract de cesiune exclusiv\u0103",
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_normal(doc,
        "\u2610 Sunt crea\u021bii originale ale Autorului",
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

    add_normal(doc,
        "Prezenta Anex\u0103 poate fi completat\u0103 prin Acte Adi\u021bionale semnate de "
        "ambele p\u0103r\u021bi, pentru opere muzicale noi livrate ulterior.",
        space_after=16)

    # Signatures for Anexa
    add_paragraph(doc, "", space_after=6)

    sig_table2 = doc.add_table(rows=4, cols=2)
    sig_table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    sig2_left = [
        ("CESIONARUL (EGI):", True),
        ("Semn\u0103tura: _______________", False),
        ("\u0218tampila:", False),
        ("Data: ____________________", False),
    ]
    sig2_right = [
        ("CEDENTUL (AUTORUL):", True),
        ("Semn\u0103tura: ____________", False),
        ("", False),
        ("Data: _________________", False),
    ]

    for i, ((left_text, left_bold), (right_text, right_bold)) in enumerate(zip(sig2_left, sig2_right)):
        cell_left = sig_table2.cell(i, 0)
        cell_left.text = ""
        p = cell_left.paragraphs[0]
        run = p.add_run(left_text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = left_bold

        cell_right = sig_table2.cell(i, 1)
        cell_right.text = ""
        p = cell_right.paragraphs[0]
        run = p.add_run(right_text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = right_bold

    # Save document
    doc.save(output_file)
    print(f"Document saved successfully to: {output_file}")
    print(f"File size: {os.path.getsize(output_file)} bytes")


if __name__ == "__main__":
    main()
