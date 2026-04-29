"""
Script pentru generarea celor 3 notificari oficiale catre
UCMR-ADA, CREDIDAM si UPFR.
Ruleaza: python generate_notifications.py
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

ORGANISMS = [
    {
        'short': 'UCMR-ADA',
        'full': 'Uniunea Compozitorilor si Muzicologilor din Romania - Asociatia pentru Drepturi de Autor (UCMR-ADA)',
        'address': 'Str. Nicolae Golescu nr. 14, Sector 1, Bucuresti',
        'rights_type': 'drepturilor patrimoniale de autor',
        'filename': 'notificare_UCMR-ADA.docx',
    },
    {
        'short': 'CREDIDAM',
        'full': 'Centrul Roman pentru Administrarea Drepturilor Artistilor Interpreti (CREDIDAM)',
        'address': 'Str. Jules Michelet nr. 15-17, Sector 1, Bucuresti',
        'rights_type': 'drepturilor conexe ale artistilor interpreti',
        'filename': 'notificare_CREDIDAM.docx',
    },
    {
        'short': 'UPFR',
        'full': 'Uniunea Producatorilor de Fonograme din Romania (UPFR)',
        'address': 'Str. Dem. I. Dobrescu nr. 4-6, Sector 1, Bucuresti',
        'rights_type': 'drepturilor conexe ale producatorilor de fonograme',
        'filename': 'notificare_UPFR.docx',
    },
]


def create_notification(org):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ── ANTET ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('SOUNDFREE S.R.L.')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0xA8, 0x6B)

    details = [
        'CUI: 54416770',
        'Nr. Reg. Comert: J2026022358004',
        'Sediu social: Iasi, Romania',
        'Email: office@soundfree.ro',
        'Telefon: 0733 272 263',
        'Web: www.soundfree.ro',
    ]
    for detail in details:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(detail)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Separator line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('_' * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x00, 0xA8, 0x6B)

    # ── DATA ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Iasi, 10 aprilie 2026')
    run.font.size = Pt(11)

    # ── DESTINATAR ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('Catre:')
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(org['full'])
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(org['address'])
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── TITLU ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run('NOTIFICARE')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0xA8, 0x6B)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(
        'privind exercitarea gestiunii individuale a '
        f'{org["rights_type"]} '
        'conform art. 145 din Legea nr. 8/1996'
    )
    run.bold = True
    run.font.size = Pt(11)

    # ── CORP ──

    # Paragraf 1 - Prezentare
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        'Subscrisa, SOUNDFREE S.R.L., cu sediul in Iasi, Romania, inregistrata la Oficiul '
        'Registrului Comertului sub nr. J2026022358004, CUI 54416770, prin reprezentant legal, '
        'in calitate de titular unic al drepturilor de autor si al drepturilor conexe asupra '
        'repertoriului muzical propriu, compus, produs si inregistrat integral de echipa SoundFree, '
        'va comunicam urmatoarele:'
    )
    run.font.size = Pt(11)

    # Paragraf 2 - Notificare formala
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('1. ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        'In conformitate cu prevederile art. 145 alin. (1) din Legea nr. 8/1996 privind dreptul '
        'de autor si drepturile conexe, republicata, cu modificarile si completarile ulterioare, '
        'va notificam in mod formal ca SOUNDFREE S.R.L. exercita '
    )
    run.font.size = Pt(11)
    run = p.add_run('gestiunea individuala ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        f'a {org["rights_type"]} asupra intregului repertoriu muzical propriu. '
        'Repertoriul nostru este compus exclusiv din opere muzicale create, produse si inregistrate '
        'integral de echipa SoundFree, fara implicarea vreunui tert titular de drepturi.'
    )
    run.font.size = Pt(11)

    # Paragraf 3 - Solicitare excludere
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('2. ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        f'Va solicitam in mod expres ca {org["short"]} sa '
    )
    run.font.size = Pt(11)
    run = p.add_run('nu includa ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        'repertoriul SoundFree in gestiunea colectiva extinsa si sa '
    )
    run.font.size = Pt(11)
    run = p.add_run('nu colecteze remuneratii ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        'in numele SOUNDFREE S.R.L. sau pentru utilizarea repertoriului nostru, '
        'intrucat ne exercitam drepturile in mod individual, direct cu utilizatorii finali.'
    )
    run.font.size = Pt(11)

    # Paragraf 4 - Licentierea directa
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('3. ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        'SOUNDFREE S.R.L. licentiaza direct utilizatorii finali (spatii comerciale — restaurante, '
        'baruri, hoteluri, magazine, birouri, saloane, cabinete etc.) prin intermediul licentelor '
        'muzicale individuale emise de la adresa www.soundfree.ro. Orice colectare de remuneratii '
        f'efectuata de {org["short"]} pentru repertoriul SoundFree este '
    )
    run.font.size = Pt(11)
    run = p.add_run('neautorizata si lipsita de temei legal')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        ', intrucat nu am mandatat nicio organizatie de gestiune colectiva sa ne reprezinte interesele.'
    )
    run.font.size = Pt(11)

    # Paragraf 5 - Consecinte juridice
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('4. ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        f'In cazul in care {org["short"]} va continua sa colecteze remuneratii sau sa pretinda '
        'drepturi asupra repertoriului SoundFree dupa primirea prezentei notificari, '
        'ne rezervam dreptul de a actiona in instanta pentru apararea drepturilor noastre '
        'patrimoniale si pentru recuperarea oricaror prejudicii cauzate.'
    )
    run.font.size = Pt(11)

    # Paragraf 6 - Repertoriu prezent si viitor
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('5. ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        'Prezenta notificare se aplica asupra '
    )
    run.font.size = Pt(11)
    run = p.add_run('intregului repertoriu muzical SoundFree, existent si viitor')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        ', incluzand toate operele muzicale publicate pe platforma www.soundfree.ro la data prezentei '
        'notificari, precum si toate operele muzicale care vor fi create, produse si adaugate ulterior '
        'in catalogul SoundFree. Catalogul complet si actualizat permanent este disponibil public '
        'la adresa www.soundfree.ro.'
    )
    run.font.size = Pt(11)

    # Paragraf 7 - Confirmare primire
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('6. ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        'Va rugam sa confirmati primirea prezentei notificari si sa luati masurile necesare '
        'pentru excluderea repertoriului SoundFree — prezent si viitor — din orice baza de date, '
        'sistem de colectare sau repartizare de remuneratii gestionat de organismul dumneavoastra.'
    )
    run.font.size = Pt(11)

    # Paragraf 8 - Anexe
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Anexe:')
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(
        '  - Certificat constatator emis de Oficiul Registrului Comertului (copie)'
    )
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(
        '  - Model de licenta muzicala SoundFree emisa catre utilizatorii finali'
    )
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(
        '  - Adresa catalogului online: www.soundfree.ro (repertoriu actualizat permanent)'
    )
    run.font.size = Pt(10)

    # ── SEMNATURA ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run('Cu stima,')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('SOUNDFREE S.R.L.')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0xA8, 0x6B)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('Prin reprezentant legal')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('_________________________')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    p = doc.add_paragraph()
    run = p.add_run('(nume, prenume, semnatura, stampila)')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # ── FOOTER TEXT ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Prezenta notificare a fost intocmita in 2 (doua) exemplare originale, '
        'din care unul pentru destinatar si unul pentru emitent.'
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # Save
    output_dir = os.path.join(os.path.dirname(__file__), 'static')
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, org['filename'])
    doc.save(filepath)
    print(f'Creat: {filepath}')


if __name__ == '__main__':
    for org in ORGANISMS:
        create_notification(org)
    print('\nToate cele 3 notificari au fost generate in static/')
